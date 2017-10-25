from docx import Document
from googletrans import Translator


def translate_doc(filename, destination='zh-CN', mix=True):
    """
    translate a word document type of file and save the result as document and keep the exactly same file format. 
        :param filename: word doc file 
        :param destination='zh-CN': 
        :param mix=True: if True, will have original language and target language into the same doc. paragraphs by paragraphs.
    """
    def tx(t): return Translator().translate(t, dest=destination).text
    doc = Document(filename)
    for p in doc.paragraphs:
        txd = tx(p.text)

        p.text = p.text + ('\n' + txd if mix else '')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txd = tx(cell.text)
                p.text = cell.text + ('\n' + txd if mix else '')

    f = filename.replace('.doc', destination.lower() + '.doc')
    doc.save(f)

if __name__ == '__main__':
    filename = 'p1.docx'
    translate_doc(filename)
